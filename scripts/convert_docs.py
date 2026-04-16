import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF

# Paths
SOURCE_MD = "docs/MASTER_PROJECT_DOCUMENTATION.md"
OUTPUT_DOCX = "docs/MASTER_PROJECT_DOCUMENTATION.docx"
OUTPUT_PDF = "docs/MASTER_PROJECT_DOCUMENTATION.pdf"

class ProfessionalPDF(FPDF):
    def header(self):
        if self.page_no() > 1: # Don't show header on title page if we had one, but here we just show it everywhere
            self.set_font('Helvetica', 'B', 10)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, 'Plagiarism Detection Service - Technical Documentation', 0, 1, 'R')
            self.line(10, 18, 200, 18)
            self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('Helvetica', 'I', 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def clean_text(text):
    # Remove v3.0.0 and markdown bold markers
    text = text.replace("(v3.0.0)", "").strip()
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Remove link markers [text](#link) -> text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text.encode('latin-1', 'replace').decode('latin-1')

def convert_to_docx(content):
    doc = Document()
    
    # Remove v3.0.0 from title
    content = content.replace("(v3.0.0)", "")
    
    lines = content.splitlines()
    is_toc = False
    
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(clean_text(line[2:]), level=0)
        elif line.startswith('## '):
            text = clean_text(line[3:])
            doc.add_heading(text, level=1)
            is_toc = "Table of Contents" in text
        elif line.startswith('### '):
            doc.add_heading(clean_text(line[4:]), level=2)
        elif line.startswith('* '):
            p = doc.add_paragraph(clean_text(line[2:]), style='List Bullet')
        elif line.startswith('    * '): # Sub-list
            p = doc.add_paragraph(clean_text(line[6:]), style='List Bullet 2')
        elif line.startswith('1. ') or re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(clean_text(text), style='List Number')
            if is_toc: # Indent TOC items
                 p.paragraph_format.left_indent = Pt(20)
        elif line.strip() == '---':
            doc.add_page_break()
        elif line.startswith('```'):
            continue 
        elif line.strip():
            doc.add_paragraph(clean_text(line))
            
    doc.save(OUTPUT_DOCX)
    print(f"Successfully created {OUTPUT_DOCX}")

def convert_to_pdf(content):
    pdf = ProfessionalPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title Styling
    pdf.set_font("Helvetica", 'B', 24)
    pdf.set_text_color(20, 50, 100)
    pdf.cell(0, 30, "Project Documentation", ln=True, align='C')
    pdf.set_font("Helvetica", 'B', 18)
    pdf.cell(0, 10, "Plagiarism Detection Service", ln=True, align='C')
    pdf.ln(20)
    
    lines = content.splitlines()
    is_toc = False
    
    for line in lines:
        if line.startswith('# '):
            continue # Already handled in title
        elif line.startswith('## '):
            pdf.ln(10)
            text = clean_text(line[3:])
            is_toc = "Table of Contents" in text
            pdf.set_font("Helvetica", 'B', 16)
            pdf.set_text_color(30, 70, 120)
            pdf.cell(0, 12, text, ln=True)
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(4)
        elif line.startswith('### '):
            pdf.ln(5)
            pdf.set_font("Helvetica", 'B', 13)
            pdf.set_text_color(50, 50, 50)
            pdf.cell(0, 10, clean_text(line[4:]), ln=True)
        elif line.startswith('    * '): # Sub-bullet in TOC
            pdf.set_font("Helvetica", size=10)
            pdf.set_text_color(80, 80, 80)
            pdf.set_x(25)
            pdf.cell(0, 7, f"- {clean_text(line[6:])}", ln=True)
        elif line.startswith('* '):
            pdf.set_font("Helvetica", size=11)
            pdf.set_text_color(0, 0, 0)
            pdf.set_x(15)
            pdf.multi_cell(0, 7, f"\x07  {clean_text(line[2:])}")
        elif line.startswith('1. ') or re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            pdf.set_font("Helvetica", size=11)
            pdf.set_text_color(0, 0, 0)
            if is_toc:
                pdf.set_x(20)
                pdf.cell(0, 8, f"{line.split('.')[0]}. {clean_text(text)}", ln=True)
            else:
                pdf.set_x(15)
                pdf.multi_cell(0, 8, f"{line.split('.')[0]}. {clean_text(text)}")
        elif line.strip() == '---':
            pdf.add_page()
        elif line.startswith('```'):
            continue
        elif line.strip():
            pdf.set_font("Helvetica", size=11)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 6, clean_text(line))
            pdf.ln(2)
            
    pdf.output(OUTPUT_PDF)
    print(f"Successfully created {OUTPUT_PDF}")

def main():
    if not os.path.exists(SOURCE_MD):
        print(f"Error: {SOURCE_MD} not found.")
        return

    with open(SOURCE_MD, "r", encoding="utf-8") as f:
        content = f.read()

    print("Generating Enhanced DOCX...")
    convert_to_docx(content)
    
    print("Generating Professional PDF...")
    convert_to_pdf(content)

if __name__ == "__main__":
    main()
